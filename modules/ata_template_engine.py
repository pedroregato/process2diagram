# modules/ata_template_engine.py
# ─────────────────────────────────────────────────────────────────────────────
# Extração de modelo (template) de ata a partir de um .docx de referência, e
# aplicação desse modelo na geração de uma ata via modules/minutes_exporter.py.
#
# PC160 (melhorias/templates-ata-por-contexto.md) — Fase 1: modelo declarativo.
# Não é um BaseAgent (sem chamada de LLM) — pura lógica Python determinística,
# mesmo espírito de agent_mermaid.py/agent_validator.py.
#
# Public API:
#   extract_template_from_docx(docx_bytes) -> (template_markdown, style_spec, assets)
#   apply_template_to_docx(minutes, style_spec, assets) -> bytes
# ─────────────────────────────────────────────────────────────────────────────

from __future__ import annotations

from io import BytesIO
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from core.knowledge_hub import MinutesModel

# Estilos de heading em Word — nomes variam por idioma/template (EN "Heading N",
# PT-BR "Título N"). Mapeamos ambos para o mesmo nível markdown.
_HEADING_STYLE_LEVELS = {
    "heading 1": 1, "título 1": 1, "titulo 1": 1,
    "heading 2": 2, "título 2": 2, "titulo 2": 2,
    "heading 3": 3, "título 3": 3, "titulo 3": 3,
    "heading 4": 4, "título 4": 4, "titulo 4": 4,
}


def _heading_level(paragraph) -> int | None:
    style_name = (paragraph.style.name or "").strip().lower()
    return _HEADING_STYLE_LEVELS.get(style_name)


def _is_pseudo_heading_run(run) -> bool:
    color = getattr(run.font, "color", None)
    rgb = getattr(color, "rgb", None) if color else None
    return bool(run.bold) and rgb is not None


def _collect_heading_paragraphs(doc) -> list[tuple[object, int]]:
    """
    (paragraph, level) pairs identifying the document's headings.

    Prefers Word's built-in Heading N / Título N paragraph styles. Many
    real-world reference docs (found in practice — corporate templates
    built in plain "Normal"/"List Paragraph" styles with manual bold +
    color formatting standing in for headings) never use those styles at
    all — for those, falls back to paragraphs whose first run is bold with
    an explicit font color, ranking distinct font sizes (descending) into
    levels 1-4. The fallback only kicks in when zero style-based headings
    exist in the whole document, so a template that DOES use real Heading
    styles is never second-guessed by this heuristic.
    """
    style_based = [
        (p, lvl) for p in doc.paragraphs
        if (lvl := _heading_level(p)) is not None and p.text.strip()
    ]
    if style_based:
        return style_based

    candidates: list[tuple[object, float | None]] = []
    for p in doc.paragraphs:
        if not p.text.strip() or not p.runs:
            continue
        if _is_pseudo_heading_run(p.runs[0]):
            size = p.runs[0].font.size
            candidates.append((p, size.pt if size else None))
    if not candidates:
        return []

    sizes = sorted({s for _, s in candidates if s is not None}, reverse=True)
    size_to_level = {s: i + 1 for i, s in enumerate(sizes[:4])}
    fallback_level = min(len(sizes), 4) + 1 if len(sizes) < 4 else 4
    return [(p, size_to_level.get(size, fallback_level)) for p, size in candidates]


def _dominant_heading_color(doc, heading_paragraphs: list[tuple[object, int]] | None = None) -> str | None:
    """
    Best-effort accent color: the most common explicit RGB color set on a
    run within a heading paragraph (as identified by
    _collect_heading_paragraphs). Returns a '#RRGGBB' string, or None if no
    heading run has an explicit color (falls back to the app default).
    """
    from collections import Counter

    if heading_paragraphs is None:
        heading_paragraphs = _collect_heading_paragraphs(doc)

    counts: Counter[str] = Counter()
    for p, _level in heading_paragraphs:
        for run in p.runs:
            color = getattr(run.font, "color", None)
            rgb = getattr(color, "rgb", None) if color else None
            if rgb:
                counts[str(rgb)] += 1
    if not counts:
        return None
    return "#" + counts.most_common(1)[0][0]


def _iter_body_blocks(doc) -> list[tuple[str, object]]:
    """
    ('paragraph' | 'table', wrapped python-docx object) pairs in real
    document order. doc.paragraphs and doc.tables are two SEPARATE lists
    with no relative order between them — this is the only way to know
    "this table came right after that heading".
    """
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    blocks: list[tuple[str, object]] = []
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            blocks.append(("paragraph", Paragraph(child, doc)))
        elif child.tag == qn("w:tbl"):
            blocks.append(("table", Table(child, doc)))
    return blocks


def _is_list_paragraph(paragraph) -> bool:
    style_name = (paragraph.style.name or "").strip().lower()
    if style_name in ("list bullet", "list number", "list paragraph"):
        return True
    pPr = paragraph._p.pPr
    return pPr is not None and pPr.numPr is not None


def _table_header_texts(table) -> list[str]:
    """Cell texts of the table's first row, stripped — best-effort column
    names, whether or not that row is visually 'header-styled'."""
    if not table.rows:
        return []
    return [cell.text.strip() for cell in table.rows[0].cells]


def _extract_sections(doc, heading_paragraphs: list[tuple[object, int]]) -> list[dict]:
    """
    One entry per heading, in document order:
      {"heading_text": str, "level": int,
       "content_kind": "table" | "list" | "paragraph" | "empty",
       "columns": list[str] | None}

    content_kind/columns are derived from every block found between this
    heading and the next heading (blank spacer paragraphs don't count):
    a table anywhere in the section wins (uses the FIRST table's header
    row — matches the common case of a single table per section, and
    avoids misclassifying a section as "paragraph" just because it opens
    with an introductory sentence before the table); otherwise a
    bullet/numbered-list paragraph makes it "list"; otherwise any other
    non-blank paragraph makes it "paragraph". A table with no readable
    header row (blank/ambiguous first row) reports columns=[] — callers
    degrade to list rendering in that case rather than inventing column
    names.
    """
    if not heading_paragraphs:
        return []

    heading_ids = {id(p._p) for p, _ in heading_paragraphs}
    level_by_id = {id(p._p): lvl for p, lvl in heading_paragraphs}
    text_by_id = {id(p._p): p.text.strip() for p, _ in heading_paragraphs}

    blocks = _iter_body_blocks(doc)

    sections: list[dict] = []
    current: dict | None = None
    section_blocks: list[tuple[str, object]] = []

    def _finalize(section: dict, s_blocks: list[tuple[str, object]]) -> dict:
        table = next((obj for kind, obj in s_blocks if kind == "table"), None)
        if table is not None:
            columns = _table_header_texts(table)
            section["content_kind"] = "table"
            section["columns"] = columns if any(columns) else []
            return section

        list_p = next(
            (obj for kind, obj in s_blocks
             if kind == "paragraph" and obj.text.strip() and _is_list_paragraph(obj)),
            None,
        )
        if list_p is not None:
            section["content_kind"] = "list"
            return section

        para = next(
            (obj for kind, obj in s_blocks if kind == "paragraph" and obj.text.strip()),
            None,
        )
        if para is not None:
            section["content_kind"] = "paragraph"
            return section

        section["content_kind"] = "empty"
        return section

    for kind, obj in blocks:
        if kind == "paragraph" and id(obj._p) in heading_ids:
            if current is not None:
                sections.append(_finalize(current, section_blocks))
            current = {
                "heading_text": text_by_id[id(obj._p)],
                "level": level_by_id[id(obj._p)],
                "content_kind": "empty",
                "columns": None,
            }
            section_blocks = []
            continue

        if current is None:
            continue
        section_blocks.append((kind, obj))

    if current is not None:
        sections.append(_finalize(current, section_blocks))

    return sections


def _extract_images_from_part(part, origin: str) -> list[dict]:
    """Pull every image relationship out of a document/header/footer part."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    assets: list[dict] = []
    try:
        rels = part.rels
    except Exception:
        return assets

    for rel in rels.values():
        if rel.reltype != RT.IMAGE or rel.is_external:
            continue
        try:
            image_part = rel.target_part
            blob = image_part.blob
            content_type = image_part.content_type or "image/png"
            try:
                width_px, height_px = image_part.image.size
            except Exception:
                width_px = height_px = None
        except Exception:
            continue

        asset_type = "logo" if origin == "header" else (
            "footer_image" if origin == "footer" else "other"
        )
        assets.append({
            "asset_type": asset_type,
            "origin": origin,
            "image_bytes": blob,
            "mime_type": content_type,
            "width_px": width_px,
            "height_px": height_px,
        })
    return assets


def _extract_page_background(doc) -> dict | None:
    """
    Best-effort extraction of a legacy VML page background (w:background in
    document.xml, referencing a v:background picture fill via r:id). This is
    an old Word feature with no high-level python-docx API — if the document
    doesn't use it, or the XML shape doesn't match what we expect, this
    silently returns None rather than failing the whole extraction.
    """
    try:
        ns = {
            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
            "v": "urn:schemas-microsoft-com:vml",
            "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        }
        root = doc.element
        bg = root.find(".//w:background", ns)
        if bg is None:
            return None
        fill = bg.find(".//v:fill", ns) if bg is not None else None
        if fill is None:
            return None
        r_id = fill.get(f"{{{ns['r']}}}id")
        if not r_id:
            return None
        part = doc.part.related_parts.get(r_id)
        if part is None:
            return None
        blob = part.blob
        content_type = getattr(part, "content_type", "image/png")
        return {
            "asset_type": "background",
            "origin": "page_background",
            "image_bytes": blob,
            "mime_type": content_type,
            "width_px": None,
            "height_px": None,
        }
    except Exception:
        return None


def extract_template_from_docx(docx_bytes: bytes) -> tuple[str, dict, list[dict]]:
    """
    Parse a reference .docx and derive:
      - template_markdown: a heading skeleton mirroring the document's own
        section names/order ('# Título', '## Subtítulo', ...).
      - style_spec: {"accent_color": "#RRGGBB" | None,
        "sections": [{"heading_text", "level", "content_kind": "table" |
        "list" | "paragraph" | "empty", "columns": list[str] | None}, ...]}
        — per-section shape (table/list/paragraph), so the export side can
        follow the reference doc's own visual form (PC160 Fase 2) instead
        of a single fixed layout. Never assumes any particular section
        name/column set — works for any template shape.
      - assets: list of {asset_type, origin, image_bytes, mime_type,
        width_px, height_px} — every image found in header/footer/body,
        plus a best-effort page background.

    Never raises for a well-formed .docx; malformed input propagates as a
    python-docx exception for the caller to surface to the user.
    """
    from docx import Document

    doc = Document(BytesIO(docx_bytes))

    # ── Markdown skeleton from heading paragraphs ──────────────────────────
    heading_paragraphs = _collect_heading_paragraphs(doc)
    lines: list[str] = []
    for p, level in heading_paragraphs:
        text = p.text.strip()
        if text:
            lines.append("#" * min(level, 4) + " " + text)
    template_markdown = "\n\n".join(lines)

    # ── Style spec ───────────────────────────────────────────────────────
    style_spec = {
        "accent_color": _dominant_heading_color(doc, heading_paragraphs),
        "sections": _extract_sections(doc, heading_paragraphs),
    }

    # ── Assets: header/footer (per section) + body + page background ──────
    assets: list[dict] = []
    seen_blobs: set[bytes] = set()

    for section in doc.sections:
        for part, origin in (
            (section.header.part, "header"),
            (section.footer.part, "footer"),
        ):
            for asset in _extract_images_from_part(part, origin):
                if asset["image_bytes"] in seen_blobs:
                    continue
                seen_blobs.add(asset["image_bytes"])
                assets.append(asset)

    for asset in _extract_images_from_part(doc.part, "body"):
        if asset["image_bytes"] in seen_blobs:
            continue
        seen_blobs.add(asset["image_bytes"])
        assets.append(asset)

    bg_asset = _extract_page_background(doc)
    if bg_asset and bg_asset["image_bytes"] not in seen_blobs:
        assets.append(bg_asset)

    return template_markdown, style_spec, assets


def apply_template_to_docx(
    minutes: "MinutesModel",
    style_spec: dict | None = None,
    assets: list[dict] | None = None,
) -> bytes:
    """
    Thin wrapper over modules.minutes_exporter.to_docx() that forwards a
    template spec built from extract_template_from_docx() (or loaded back
    from the ata_templates/ata_template_assets tables). Kept as a separate
    function (rather than requiring every caller to import to_docx directly)
    so the template-resolution/application concern stays in one module.
    """
    from modules.minutes_exporter import to_docx

    template_spec = None
    if style_spec or assets:
        template_spec = {
            "accent_color": (style_spec or {}).get("accent_color"),
            "assets": assets or [],
            "sections": (style_spec or {}).get("sections") or [],
        }
    return to_docx(minutes, template_spec=template_spec)
