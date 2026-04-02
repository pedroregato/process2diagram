# modules/ingest.py (versão completa para .txt, .docx, .pdf)
import io
from pathlib import Path

def load_transcript(uploaded_file) -> str:
    filename = uploaded_file.name.lower()
    
    if filename.endswith('.txt'):
        try:
            return uploaded_file.read().decode('utf-8')
        except UnicodeDecodeError:
            uploaded_file.seek(0)
            return uploaded_file.read().decode('latin-1')
    
    elif filename.endswith('.docx'):
        try:
            import docx
            doc = docx.Document(uploaded_file)
            full_text = []
            # Parágrafos
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            # Tabelas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if para.text.strip():
                                full_text.append(para.text)
            # Cabeçalhos e rodapés (opcional)
            for section in doc.sections:
                for para in section.header.paragraphs:
                    if para.text.strip():
                        full_text.append(para.text)
                for para in section.footer.paragraphs:
                    if para.text.strip():
                        full_text.append(para.text)
            return '\n'.join(full_text)
        except ImportError:
            raise ImportError("python-docx is required. Install: pip install python-docx")
        except Exception as e:
            raise RuntimeError(f"Failed to parse .docx: {e}")
    
    elif filename.endswith('.pdf'):
        try:
            from pypdf import PdfReader
        except ImportError:
            try:
                from PyPDF2 import PdfReader
            except ImportError:
                raise ImportError("pypdf or PyPDF2 required. Install: pip install pypdf")
        reader = PdfReader(uploaded_file)
        text = '\n'.join(page.extract_text() or '' for page in reader.pages)
        return text
    
    else:
        raise ValueError(f"Unsupported file type: {filename}")
