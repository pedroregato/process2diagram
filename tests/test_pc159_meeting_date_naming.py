# tests/test_pc159_meeting_date_naming.py
"""
Tests for PC159:
- services/export_service.py::format_date_suffix() — normalizes a meeting
  date (Brazilian DD/MM/AAAA from LLM extraction, ISO from the DB, or a
  datetime.date object) into a filename-safe AAAA-MM-DD suffix, so exported
  artifacts are named after the MEETING date instead of the download date.
- modules/minutes_exporter.py::to_docx() markdown fallback — a meeting
  loaded from the DB with only minutes_md (no structured fields) used to
  produce a near-empty Word document; it must now render the markdown
  content instead.
"""

from datetime import date
from io import BytesIO

from services.export_service import format_date_suffix, make_filename
from core.knowledge_hub import MinutesModel
from modules.minutes_exporter import to_docx


class TestFormatDateSuffix:
    def test_brazilian_format(self):
        assert format_date_suffix("05/01/2026") == "2026-01-05"

    def test_iso_date_only(self):
        assert format_date_suffix("2026-01-05") == "2026-01-05"

    def test_iso_timestamp_truncated(self):
        assert format_date_suffix("2026-01-05T14:30:00") == "2026-01-05"

    def test_date_object(self):
        assert format_date_suffix(date(2026, 1, 5)) == "2026-01-05"

    def test_none_falls_back_to_today(self):
        assert format_date_suffix(None) == date.today().isoformat()

    def test_empty_string_falls_back_to_today(self):
        assert format_date_suffix("") == date.today().isoformat()

    def test_placeholder_dash_falls_back_to_today(self):
        assert format_date_suffix("—") == date.today().isoformat()

    def test_used_in_make_filename(self):
        name = make_filename("minutes", "html", "P2D_", format_date_suffix("05/01/2026"))
        assert name == "P2D_minutes2026-01-05.html"


class TestToDocxMarkdownFallback:
    def test_markdown_only_minutes_produces_real_content(self):
        m = MinutesModel(title="Reunião Antiga", date="2026-01-05", ready=True)
        m.minutes_md = (
            "# Ata\n\nTexto de abertura.\n\n"
            "## Decisões\n- Decisão 1\n- Decisão 2\n\n"
            "| Col A | Col B |\n|---|---|\n| valor 1 | valor 2 |\n"
        )
        data = to_docx(m)
        assert len(data) > 5000  # a near-empty stub is ~a few KB smaller

        from docx import Document
        doc = Document(BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        assert "Ata" in paragraphs
        assert "Texto de abertura." in paragraphs
        assert "Decisões" in paragraphs
        assert "Decisão 1" in paragraphs
        assert "Decisão 2" in paragraphs
        assert len(doc.tables) == 1
        rows = [[c.text for c in row.cells] for row in doc.tables[0].rows]
        assert rows == [["Col A", "Col B"], ["valor 1", "valor 2"]]

    def test_structured_minutes_do_not_use_fallback(self):
        """Regression guard: a minutes with real structured fields must still
        render via the normal section builders, not the markdown fallback,
        even if minutes_md also happens to be populated."""
        from core.knowledge_hub import ActionItem
        m = MinutesModel(
            title="Reunião Estruturada", date="2026-01-05",
            decisions=["Decisão estruturada"],
            action_items=[ActionItem(task="Fazer X", responsible="Pedro")],
            minutes_md="# Isto não deve aparecer\n",
            ready=True,
        )
        data = to_docx(m)
        from docx import Document
        doc = Document(BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        assert "Decisão estruturada" in paragraphs
        assert "Isto não deve aparecer" not in paragraphs

    def test_empty_minutes_md_produces_minimal_doc_without_error(self):
        m = MinutesModel(title="Reunião Vazia", date="2026-01-05", ready=True)
        data = to_docx(m)
        assert isinstance(data, bytes)
        assert len(data) > 0
