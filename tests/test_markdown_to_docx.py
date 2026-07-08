# tests/test_markdown_to_docx.py
"""
Tests for modules/minutes_exporter.py::markdown_to_docx() — generic
Markdown -> .docx converter used by export_project_charter_docx() and any
future LLM-generated Markdown artifact that isn't a MinutesModel.

No real DB/LLM calls.
"""

from io import BytesIO

from docx import Document

from modules.minutes_exporter import markdown_to_docx


class TestMarkdownToDocx:
    def test_produces_valid_docx_bytes(self):
        data = markdown_to_docx("# Título\n\nTexto simples.")
        assert isinstance(data, bytes) and len(data) > 0
        Document(BytesIO(data))  # must not raise

    def test_headings_and_bullets_render_as_real_paragraphs(self):
        md = "# Documento\n\n## Seção 1\n\n- Item A\n- Item B\n\n## Seção 2\n\nParágrafo comum."
        data = markdown_to_docx(md)
        doc = Document(BytesIO(data))
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert "Documento" in texts
        assert "Seção 1" in texts
        assert "Item A" in texts
        assert "Item B" in texts
        assert "Seção 2" in texts
        assert "Parágrafo comum." in texts

    def test_markdown_table_becomes_real_docx_table(self):
        md = (
            "# Charter\n\n"
            "## Riscos\n\n"
            "| Risco | Severidade |\n"
            "|---|---|\n"
            "| Atraso | Alta |\n"
        )
        data = markdown_to_docx(md)
        doc = Document(BytesIO(data))
        assert len(doc.tables) == 1
        table = doc.tables[0]
        assert [c.text for c in table.rows[0].cells] == ["Risco", "Severidade"]
        assert [c.text for c in table.rows[1].cells] == ["Atraso", "Alta"]

    def test_empty_markdown_still_produces_valid_docx(self):
        data = markdown_to_docx("")
        assert isinstance(data, bytes) and len(data) > 0
        Document(BytesIO(data))  # must not raise
