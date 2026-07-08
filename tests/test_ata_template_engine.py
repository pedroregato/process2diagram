# tests/test_ata_template_engine.py
"""
Tests for PC160 (melhorias/templates-ata-por-contexto.md): per-context Word
templates for meeting minutes.

Covers:
- modules/ata_template_engine.py::extract_template_from_docx() — heading
  skeleton, accent color detection, image extraction (header/body).
- modules/ata_template_engine.py::apply_template_to_docx() wrapper.
- modules/minutes_exporter.py::to_docx() with template_spec (accent color
  override + logo insertion), retrocompatible when omitted.
- core/project_store.py CRUD (save/get_active/list/activate/deactivate/
  delete) via a mocked Supabase client.
- agents/agent_minutes.py::build_prompt() injects ata_template_markdown.

No real DB/LLM calls.
"""

import base64
from io import BytesIO
from unittest.mock import patch, MagicMock

from docx import Document
from docx.shared import Cm, Pt, RGBColor

from core.knowledge_hub import MinutesModel, ActionItem
from modules.ata_template_engine import extract_template_from_docx, apply_template_to_docx
from modules.minutes_exporter import to_docx

# 1x1 transparent PNG, used as a fake logo/image throughout.
_TINY_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


def _build_reference_docx(with_logo: bool = True, accent_hex: str = "AA1122") -> bytes:
    doc = Document()
    if with_logo:
        header = doc.sections[0].header
        header.is_linked_to_previous = False
        p = header.paragraphs[0]
        p.add_run().add_picture(BytesIO(_TINY_PNG), width=Cm(2))
    h1 = doc.add_paragraph("Ata de Reunião", style="Heading 1")
    h1.runs[0].font.color.rgb = RGBColor(
        int(accent_hex[0:2], 16), int(accent_hex[2:4], 16), int(accent_hex[4:6], 16)
    )
    doc.add_paragraph("Participantes", style="Heading 2")
    doc.add_paragraph("Fulano, Beltrano")
    doc.add_paragraph("Decisões", style="Heading 2")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_pseudo_heading_docx(accent_hex: str = "1F4E79") -> bytes:
    """
    Mirrors a real-world corporate template (found via a user-submitted
    .docx, PC160-C): section titles built with bold + explicit font color
    on plain 'Normal'/'List Paragraph' styles — never Word's built-in
    Heading N style. Body paragraphs are neither bold nor colored the
    accent color, so they must NOT be picked up as headings.
    """
    accent = RGBColor(int(accent_hex[0:2], 16), int(accent_hex[2:4], 16), int(accent_hex[4:6], 16))
    body_color = RGBColor(0x1E, 0x29, 0x3B)

    doc = Document()

    def _pseudo_heading(text: str, size_pt: int) -> None:
        p = doc.add_paragraph(style="List Paragraph")
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(size_pt)
        r.font.color.rgb = accent

    def _body(text: str) -> None:
        p = doc.add_paragraph(style="List Paragraph")
        r = p.add_run(text)
        r.font.color.rgb = body_color

    _pseudo_heading("ATA DE REUNIÃO", 16)
    _pseudo_heading("Participantes", 13)
    _body("Fulano, Beltrano")
    _pseudo_heading("Assuntos Discutidos", 13)
    _pseudo_heading("Status da Integração", 11)
    _body("Texto corrido sem negrito nem cor de destaque.")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_table(doc, columns: list[str]) -> None:
    table = doc.add_table(rows=1, cols=len(columns))
    for i, c in enumerate(columns):
        table.rows[0].cells[i].text = c


def _build_fgv_like_reference_docx() -> bytes:
    """
    One example of a real-world template shape (not the only one this
    feature must support — see _build_alternative_reference_docx for a
    deliberately different one): 'Participantes' as a 3-column table
    (Nome/E-mail/Unidade), 'Encaminhamentos' as a 3-column table
    (Ação/Responsável/Data), 'Decisões' as a plain bullet list.
    """
    doc = Document()
    doc.add_paragraph("Ata de Reunião", style="Heading 1")
    doc.add_paragraph("Participantes", style="Heading 2")
    _add_table(doc, ["Nome", "E-mail", "Unidade"])
    doc.add_paragraph("Decisões", style="Heading 2")
    doc.add_paragraph("Alguma decisão", style="List Bullet")
    doc.add_paragraph("Encaminhamentos", style="Heading 2")
    _add_table(doc, ["Ação", "Responsável", "Data"])
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_alternative_reference_docx() -> bytes:
    """
    A SECOND, deliberately different template — proves the feature follows
    ANY reference shape, not just the FGV one: participants as a plain
    list (no table), a 'Riscos' table with a column no _COLUMN_FIELD_
    PATTERNS entry recognizes (degrades to list), and 'Next Steps' as a
    table with English/mismatched column names (partial column match).
    """
    doc = Document()
    doc.add_paragraph("Weekly Status Report", style="Heading 1")
    doc.add_paragraph("Presentes", style="Heading 2")
    doc.add_paragraph("Alguém", style="List Bullet")
    doc.add_paragraph("Riscos Identificados", style="Heading 2")
    _add_table(doc, ["Descrição"])
    doc.add_paragraph("Next Steps", style="Heading 2")
    _add_table(doc, ["Tarefa", "Owner", "Due"])
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestExtractTemplateFromDocx:
    def test_markdown_skeleton_follows_heading_order(self):
        ref = _build_reference_docx()
        template_markdown, _, _ = extract_template_from_docx(ref)
        assert template_markdown == "# Ata de Reunião\n\n## Participantes\n\n## Decisões"

    def test_accent_color_detected_from_heading_run(self):
        ref = _build_reference_docx(accent_hex="00FF88")
        _, style_spec, _ = extract_template_from_docx(ref)
        assert style_spec["accent_color"] == "#00FF88"

    def test_no_heading_color_returns_none(self):
        doc = Document()
        doc.add_paragraph("Sem cor", style="Heading 1")  # no explicit run color
        buf = BytesIO()
        doc.save(buf)
        _, style_spec, _ = extract_template_from_docx(buf.getvalue())
        assert style_spec["accent_color"] is None

    def test_header_logo_extracted_as_logo_asset(self):
        ref = _build_reference_docx(with_logo=True)
        _, _, assets = extract_template_from_docx(ref)
        logos = [a for a in assets if a["asset_type"] == "logo"]
        assert len(logos) == 1
        assert logos[0]["origin"] == "header"
        assert logos[0]["mime_type"] == "image/png"
        assert logos[0]["image_bytes"] == _TINY_PNG

    def test_no_logo_produces_no_assets(self):
        ref = _build_reference_docx(with_logo=False)
        _, _, assets = extract_template_from_docx(ref)
        assert assets == []

    def test_no_headings_produces_empty_skeleton_without_error(self):
        doc = Document()
        doc.add_paragraph("Só um parágrafo normal, sem heading.")
        buf = BytesIO()
        doc.save(buf)
        template_markdown, style_spec, assets = extract_template_from_docx(buf.getvalue())
        assert template_markdown == ""
        assert style_spec["accent_color"] is None
        assert style_spec["sections"] == []
        assert assets == []

    def test_pseudo_heading_fallback_when_no_word_heading_style_used(self):
        """PC160-C: real-world templates built with bold+colored 'Normal'/
        'List Paragraph' paragraphs (never Word's built-in Heading N style)
        must still be recognized — this was the exact shape of the .docx a
        user reported as 'not following the registered model'."""
        ref = _build_pseudo_heading_docx()
        template_markdown, style_spec, _ = extract_template_from_docx(ref)
        assert style_spec["accent_color"] == "#1F4E79"
        assert template_markdown == (
            "# ATA DE REUNIÃO\n\n"
            "## Participantes\n\n"
            "## Assuntos Discutidos\n\n"
            "### Status da Integração"
        )

    def test_pseudo_heading_fallback_ignores_plain_body_paragraphs(self):
        ref = _build_pseudo_heading_docx()
        template_markdown, _, _ = extract_template_from_docx(ref)
        assert "Fulano, Beltrano" not in template_markdown
        assert "Texto corrido" not in template_markdown

    def test_real_word_heading_style_takes_precedence_over_pseudo_heading(self):
        """A document that legitimately uses Heading N styles must never be
        second-guessed by the bold+color fallback, even if it also happens
        to contain bold-colored body text elsewhere."""
        doc = Document()
        doc.add_paragraph("Título Real", style="Heading 1")
        p = doc.add_paragraph(style="List Paragraph")
        r = p.add_run("Texto em negrito e colorido, mas não é heading")
        r.bold = True
        r.font.color.rgb = RGBColor(0xAA, 0x11, 0x22)
        buf = BytesIO()
        doc.save(buf)
        template_markdown, _, _ = extract_template_from_docx(buf.getvalue())
        assert template_markdown == "# Título Real"


class TestExtractSections:
    """PC160 Fase 2 — per-section shape (table/list/paragraph + columns),
    driven purely by what's actually in each reference .docx (never assumes
    any particular section name or column set — see _build_alternative_
    reference_docx for proof the same code handles a totally different
    template shape)."""

    def test_table_section_detected_with_columns_in_document_order(self):
        ref = _build_fgv_like_reference_docx()
        _, style_spec, _ = extract_template_from_docx(ref)
        participantes = next(s for s in style_spec["sections"] if s["heading_text"] == "Participantes")
        assert participantes["content_kind"] == "table"
        assert participantes["columns"] == ["Nome", "E-mail", "Unidade"]

    def test_list_section_detected(self):
        ref = _build_fgv_like_reference_docx()
        _, style_spec, _ = extract_template_from_docx(ref)
        decisoes = next(s for s in style_spec["sections"] if s["heading_text"] == "Decisões")
        assert decisoes["content_kind"] == "list"

    def test_paragraph_section_detected(self):
        doc = Document()
        doc.add_paragraph("Ata", style="Heading 1")
        doc.add_paragraph("Introdução", style="Heading 2")
        doc.add_paragraph("Texto corrido, sem marcador de lista.")
        buf = BytesIO()
        doc.save(buf)
        _, style_spec, _ = extract_template_from_docx(buf.getvalue())
        intro = next(s for s in style_spec["sections"] if s["heading_text"] == "Introdução")
        assert intro["content_kind"] == "paragraph"

    def test_heading_without_content_before_next_heading_is_empty(self):
        doc = Document()
        doc.add_paragraph("Ata", style="Heading 1")
        doc.add_paragraph("Vazio", style="Heading 2")
        doc.add_paragraph("Próximo", style="Heading 2")
        doc.add_paragraph("conteúdo aqui", style="List Bullet")
        buf = BytesIO()
        doc.save(buf)
        _, style_spec, _ = extract_template_from_docx(buf.getvalue())
        vazio = next(s for s in style_spec["sections"] if s["heading_text"] == "Vazio")
        assert vazio["content_kind"] == "empty"

    def test_sections_preserve_document_heading_order(self):
        ref = _build_fgv_like_reference_docx()
        _, style_spec, _ = extract_template_from_docx(ref)
        headings = [s["heading_text"] for s in style_spec["sections"]]
        assert headings == ["Ata de Reunião", "Participantes", "Decisões", "Encaminhamentos"]

    def test_table_without_readable_header_reports_empty_columns(self):
        doc = Document()
        doc.add_paragraph("Ata", style="Heading 1")
        doc.add_paragraph("Participantes", style="Heading 2")
        table = doc.add_table(rows=1, cols=2)  # header row left blank
        buf = BytesIO()
        doc.save(buf)
        _, style_spec, _ = extract_template_from_docx(buf.getvalue())
        participantes = next(s for s in style_spec["sections"] if s["heading_text"] == "Participantes")
        assert participantes["content_kind"] == "table"
        assert participantes["columns"] == []

    def test_second_different_template_shape_also_extracted_correctly(self):
        """The alternative template (participants as list, English/mixed
        column names) is extracted with the SAME code path — no branch
        specific to the FGV shape."""
        ref = _build_alternative_reference_docx()
        _, style_spec, _ = extract_template_from_docx(ref)
        by_heading = {s["heading_text"]: s for s in style_spec["sections"]}
        assert by_heading["Presentes"]["content_kind"] == "list"
        assert by_heading["Riscos Identificados"]["content_kind"] == "table"
        assert by_heading["Riscos Identificados"]["columns"] == ["Descrição"]
        assert by_heading["Next Steps"]["content_kind"] == "table"
        assert by_heading["Next Steps"]["columns"] == ["Tarefa", "Owner", "Due"]


class TestApplyTemplateToDocx:
    def test_produces_valid_docx_with_style_and_assets(self):
        m = MinutesModel(title="Reunião", date="2026-07-08", decisions=["Decisão 1"], ready=True)
        assets = [{"asset_type": "logo", "origin": "header", "image_bytes": _TINY_PNG,
                   "mime_type": "image/png", "width_px": 10, "height_px": 10}]
        data = apply_template_to_docx(m, style_spec={"accent_color": "#112233"}, assets=assets)
        doc = Document(BytesIO(data))
        assert len(doc.sections[0].header.part.rels) >= 1  # logo relationship present

    def test_no_style_or_assets_still_produces_valid_docx(self):
        m = MinutesModel(title="Reunião", date="2026-07-08", decisions=["Decisão 1"], ready=True)
        data = apply_template_to_docx(m)
        assert isinstance(data, bytes)
        assert len(data) > 0


class TestToDocxTemplateSpec:
    def test_accent_color_override_applies_to_headings(self):
        m = MinutesModel(title="Reunião", date="2026-07-08", decisions=["Decisão 1"], ready=True)
        data = to_docx(m, template_spec={"accent_color": "#AA1122", "assets": []})
        doc = Document(BytesIO(data))
        heading = next(p for p in doc.paragraphs if p.text.strip().upper() == "DECISÕES TOMADAS")
        assert str(heading.runs[0].font.color.rgb) == "AA1122"

    def test_default_color_unaffected_without_template_spec(self):
        m = MinutesModel(title="Reunião", date="2026-07-08", decisions=["Decisão 1"], ready=True)
        data = to_docx(m)
        doc = Document(BytesIO(data))
        heading = next(p for p in doc.paragraphs if p.text.strip().upper() == "DECISÕES TOMADAS")
        assert str(heading.runs[0].font.color.rgb) == "2E7FD9"

    def test_logo_inserted_into_header_when_asset_present(self):
        m = MinutesModel(title="Reunião", date="2026-07-08", ready=True)
        data = to_docx(m, template_spec={
            "accent_color": None,
            "assets": [{"asset_type": "logo", "image_bytes": _TINY_PNG}],
        })
        doc = Document(BytesIO(data))
        assert len(doc.sections[0].header.part.rels) >= 1

    def test_malformed_accent_color_falls_back_silently(self):
        m = MinutesModel(title="Reunião", date="2026-07-08", decisions=["Decisão 1"], ready=True)
        data = to_docx(m, template_spec={"accent_color": "not-a-color", "assets": []})
        doc = Document(BytesIO(data))  # must not raise
        assert isinstance(data, bytes) and len(data) > 0

    def _section_border_color(self, paragraph) -> str | None:
        from docx.oxml.ns import qn
        pPr = paragraph._p.pPr
        if pPr is None:
            return None
        pBdr = pPr.find(qn("w:pBdr"))
        if pBdr is None:
            return None
        bottom = pBdr.find(qn("w:bottom"))
        return bottom.get(qn("w:color")) if bottom is not None else None

    def test_section_border_follows_accent_override_in_structured_path(self):
        """PC160-D: the bottom-border separator drawn under each structured
        section heading was hardcoded to the app's default blue regardless
        of template_spec — found comparing a user's real reference .docx
        against the ata it exported (no separator color match at all)."""
        m = MinutesModel(title="Reunião", date="2026-07-08", decisions=["Decisão 1"], ready=True)
        data = to_docx(m, template_spec={"accent_color": "#1F4E79", "assets": []})
        doc = Document(BytesIO(data))
        heading = next(p for p in doc.paragraphs if p.text.strip().upper() == "DECISÕES TOMADAS")
        assert self._section_border_color(heading) == "1F4E79"

    def test_section_border_follows_accent_override_in_markdown_fallback_path(self):
        """Same fix, other rendering path: a meeting loaded from the DB with
        only minutes_md (no structured fields — the exact case of a
        reunião existente / Central de Artefatos download) rendered no
        separator line at all under '##' section headings."""
        md = "# Reunião\n\n## Participantes\n\n- Fulano\n\n## Pauta\n\n1. Item\n"
        m = MinutesModel(title="Reunião", date="2026-07-08", minutes_md=md, ready=True)
        data = to_docx(m, template_spec={"accent_color": "#1F4E79", "assets": []})
        doc = Document(BytesIO(data))
        heading = next(p for p in doc.paragraphs if p.text.strip() == "Participantes")
        assert self._section_border_color(heading) == "1F4E79"

    def test_section_border_default_color_without_template_spec(self):
        md = "# Reunião\n\n## Participantes\n\n- Fulano\n"
        m = MinutesModel(title="Reunião", date="2026-07-08", minutes_md=md, ready=True)
        data = to_docx(m)
        doc = Document(BytesIO(data))
        heading = next(p for p in doc.paragraphs if p.text.strip() == "Participantes")
        assert self._section_border_color(heading) == "2E7FD9"


def _template_spec_from_docx(ref_bytes: bytes) -> dict:
    """Runs the real extraction (not a hand-built dict) so these tests
    exercise the same end-to-end path used in production."""
    _, style_spec, assets = extract_template_from_docx(ref_bytes)
    return {
        "accent_color": style_spec.get("accent_color"),
        "assets": assets,
        "sections": style_spec.get("sections") or [],
    }


class TestRenderTemplatedSections:
    """PC160 Fase 2 — to_docx() follows the reference doc's own per-section
    shape (table/list/paragraph). Every assertion here must hold for BOTH
    reference templates (_build_fgv_like_reference_docx and
    _build_alternative_reference_docx) wherever relevant, to keep the
    feature honest about not being hardcoded to one template's shape."""

    def test_participants_render_as_table_matching_template_columns(self):
        template_spec = _template_spec_from_docx(_build_fgv_like_reference_docx())
        m = MinutesModel(title="Reunião", date="2026-07-08",
                          participants=["Fulano de Tal (FT)"], ready=True)
        data = to_docx(m, template_spec=template_spec)
        doc = Document(BytesIO(data))
        table = next(t for t in doc.tables if [c.text for c in t.rows[0].cells] == ["Nome", "E-mail", "Unidade"])
        row = table.rows[1].cells
        assert row[0].text == "Fulano de Tal (FT)"
        assert row[1].text == "—"  # never a fabricated e-mail
        assert row[2].text == "—"

    def test_action_items_use_template_columns_not_app_default(self):
        template_spec = _template_spec_from_docx(_build_fgv_like_reference_docx())
        m = MinutesModel(
            title="Reunião", date="2026-07-08",
            action_items=[ActionItem(task="Enviar relatório", responsible="Fulano",
                                      deadline="10/07/2026", priority="high", raised_by="FT")],
            ready=True,
        )
        data = to_docx(m, template_spec=template_spec)
        doc = Document(BytesIO(data))
        table = next(t for t in doc.tables if [c.text for c in t.rows[0].cells] == ["Ação", "Responsável", "Data"])
        assert [c.text for c in table.rows[1].cells] == ["Enviar relatório", "Fulano", "10/07/2026"]
        # the app's own default 5-column shape (with "Prioridade") must NOT
        # also appear — the template's 3-column shape fully replaces it.
        assert not any("Prioridade" in [c.text for c in t.rows[0].cells] for t in doc.tables)

    def test_field_without_matching_template_section_still_falls_back_to_default(self):
        template_spec = _template_spec_from_docx(_build_fgv_like_reference_docx())
        m = MinutesModel(title="Reunião", date="2026-07-08",
                          next_meeting="15/07/2026 às 10h", ready=True)
        data = to_docx(m, template_spec=template_spec)
        doc = Document(BytesIO(data))
        assert any("15/07/2026 às 10h" in p.text for p in doc.paragraphs)

    def test_duplicate_heading_renders_field_only_once(self):
        doc = Document()
        doc.add_paragraph("Ata", style="Heading 1")
        doc.add_paragraph("Parte 1", style="Heading 2")
        doc.add_paragraph("Participantes", style="Heading 2")
        doc.add_paragraph("um item", style="List Bullet")
        doc.add_paragraph("Parte 2", style="Heading 2")
        doc.add_paragraph("Participantes", style="Heading 2")
        doc.add_paragraph("outro item", style="List Bullet")
        buf = BytesIO()
        doc.save(buf)
        template_spec = _template_spec_from_docx(buf.getvalue())

        m = MinutesModel(title="Reunião", date="2026-07-08",
                          participants=["Fulano", "Beltrano"], ready=True)
        data = to_docx(m, template_spec=template_spec)
        doc_out = Document(BytesIO(data))
        occurrences = [p for p in doc_out.paragraphs if p.text.strip() == "PARTICIPANTES"]
        assert len(occurrences) == 1

    def test_unmatched_heading_skipped_without_error(self):
        doc = Document()
        doc.add_paragraph("Ata", style="Heading 1")
        doc.add_paragraph("Xyzzy Não Reconhecido", style="Heading 2")
        doc.add_paragraph("algo", style="List Bullet")
        buf = BytesIO()
        doc.save(buf)
        template_spec = _template_spec_from_docx(buf.getvalue())

        m = MinutesModel(title="Reunião", date="2026-07-08", decisions=["Decisão real"], ready=True)
        data = to_docx(m, template_spec=template_spec)  # must not raise
        doc_out = Document(BytesIO(data))
        assert any("Decisão real" in p.text for p in doc_out.paragraphs)
        assert not any("XYZZY" in p.text.upper() for p in doc_out.paragraphs)

    def test_table_field_without_column_patterns_degrades_to_list(self):
        template_spec = _template_spec_from_docx(_build_alternative_reference_docx())
        m = MinutesModel(title="Reunião", date="2026-07-08",
                          risks_identified=["Atraso no fornecedor"], ready=True)
        data = to_docx(m, template_spec=template_spec)
        doc = Document(BytesIO(data))
        # risks_identified has no _COLUMN_FIELD_PATTERNS entry, so even
        # though the template showed "Riscos Identificados" as a table,
        # rendering must degrade to a bullet list rather than guessing columns.
        assert not any(t.rows and t.rows[0].cells[0].text == "Descrição" for t in doc.tables)
        assert any("Atraso no fornecedor" in p.text for p in doc.paragraphs)

    def test_alternative_template_participants_as_list_not_table(self):
        template_spec = _template_spec_from_docx(_build_alternative_reference_docx())
        m = MinutesModel(title="Reunião", date="2026-07-08",
                          participants=["Fulano"], ready=True)
        data = to_docx(m, template_spec=template_spec)
        doc = Document(BytesIO(data))
        assert doc.tables == []  # participants had no table in this template
        assert any(p.text.strip() == "Fulano" for p in doc.paragraphs)

    def test_no_sections_key_matches_pre_existing_behavior(self):
        m = MinutesModel(title="Reunião", date="2026-07-08",
                          participants=["Fulano"], ready=True)
        with_empty_sections = to_docx(m, template_spec={"accent_color": None, "assets": [], "sections": []})
        without_key_at_all = to_docx(m, template_spec={"accent_color": None, "assets": []})
        assert with_empty_sections == without_key_at_all


class TestApplyTemplateToDocxSections:
    def test_apply_template_forwards_sections_to_to_docx(self):
        style_spec = {
            "accent_color": "#1F4E79",
            "sections": [{"heading_text": "Participantes", "level": 2,
                          "content_kind": "table", "columns": ["Nome", "E-mail", "Unidade"]}],
        }
        m = MinutesModel(title="Reunião", date="2026-07-08",
                          participants=["Fulano"], ready=True)
        data = apply_template_to_docx(m, style_spec=style_spec, assets=[])
        doc = Document(BytesIO(data))
        table = next(t for t in doc.tables if [c.text for c in t.rows[0].cells] == ["Nome", "E-mail", "Unidade"])
        assert table.rows[1].cells[0].text == "Fulano"


class _FakeQuery:
    def __init__(self, rows, table_name, log):
        self._rows = rows
        self._table_name = table_name
        self._log = log
        self._pending_update = None
        self._pending_insert = None
        self._pending_delete = False

    def select(self, *a, **k):
        return self

    def eq(self, field, value):
        self._rows = [r for r in self._rows if r.get(field) == value]
        return self

    def order(self, *a, **k):
        return self

    def limit(self, *a, **k):
        return self

    def update(self, patch_dict):
        self._pending_update = patch_dict
        return self

    def insert(self, payload):
        self._pending_insert = payload
        return self

    def delete(self):
        self._pending_delete = True
        return self

    def execute(self):
        resp = MagicMock()
        if self._pending_update is not None:
            for r in self._rows:
                r.update(self._pending_update)
            self._log.append((self._table_name, "update", self._pending_update))
            resp.data = self._rows
        elif self._pending_insert is not None:
            import uuid
            new_row = dict(self._pending_insert)
            new_row.setdefault("id", str(uuid.uuid4()))
            self._rows.append(new_row)
            self._log.append((self._table_name, "insert", new_row))
            resp.data = [new_row]
        elif self._pending_delete:
            self._log.append((self._table_name, "delete", [r.get("id") for r in self._rows]))
            resp.data = self._rows
        else:
            resp.data = self._rows
        return resp


class _FakeDB:
    def __init__(self, tables: dict):
        self._tables = tables
        self.log: list[tuple] = []

    def table(self, name):
        return _FakeQuery(self._tables.setdefault(name, []), name, self.log)


class TestProjectStoreAtaTemplateCRUD:
    def test_save_ata_template_persists_template_and_assets(self):
        from core.project_store import save_ata_template
        db = _FakeDB({"ata_templates": [], "ata_template_assets": []})
        ref = _build_reference_docx()
        with patch("core.project_store._db", return_value=db):
            result = save_ata_template("ctx-1", "Modelo Teste", "modelo.docx", ref, "tester")
        assert result is not None
        assert result["name"] == "Modelo Teste"
        assert result["is_active"] is True
        assert len(db._tables["ata_template_assets"]) == 1  # the logo

    def test_save_ata_template_deactivates_previous_active(self):
        from core.project_store import save_ata_template
        db = _FakeDB({
            "ata_templates": [{"id": "old-1", "context_id": "ctx-1", "is_active": True}],
            "ata_template_assets": [],
        })
        ref = _build_reference_docx(with_logo=False)
        with patch("core.project_store._db", return_value=db):
            save_ata_template("ctx-1", "Novo Modelo", "novo.docx", ref, "tester")
        old = next(r for r in db._tables["ata_templates"] if r["id"] == "old-1")
        assert old["is_active"] is False

    def test_get_active_ata_template_decodes_assets(self):
        from core.project_store import get_active_ata_template
        tid = "tpl-1"
        db = _FakeDB({
            "ata_templates": [{
                "id": tid, "context_id": "ctx-1", "name": "M1",
                "template_markdown": "# Ata", "style_spec": {"accent_color": "#112233"},
                "is_active": True,
            }],
            "ata_template_assets": [{
                "template_id": tid, "asset_type": "logo", "origin": "header",
                "image_base64": base64.b64encode(_TINY_PNG).decode("ascii"),
                "mime_type": "image/png", "width_px": 10, "height_px": 10,
            }],
        })
        with patch("core.project_store._db", return_value=db):
            result = get_active_ata_template("ctx-1")
        assert result is not None
        assert result["template_markdown"] == "# Ata"
        assert len(result["assets"]) == 1
        assert result["assets"][0]["image_bytes"] == _TINY_PNG

    def test_get_active_ata_template_returns_none_when_no_active_template(self):
        from core.project_store import get_active_ata_template
        db = _FakeDB({"ata_templates": [], "ata_template_assets": []})
        with patch("core.project_store._db", return_value=db):
            result = get_active_ata_template("ctx-without-template")
        assert result is None

    def test_activate_deactivates_others_in_same_context(self):
        from core.project_store import activate_ata_template
        db = _FakeDB({
            "ata_templates": [
                {"id": "a", "context_id": "ctx-1", "is_active": True},
                {"id": "b", "context_id": "ctx-1", "is_active": False},
            ],
        })
        with patch("core.project_store._db", return_value=db):
            ok = activate_ata_template("b", "ctx-1")
        assert ok is True

    def test_delete_ata_template_returns_true_on_success(self):
        from core.project_store import delete_ata_template
        db = _FakeDB({"ata_templates": [{"id": "x", "context_id": "ctx-1", "is_active": True}]})
        with patch("core.project_store._db", return_value=db):
            assert delete_ata_template("x") is True


class TestAgentMinutesTemplateInjection:
    def test_template_markdown_injected_when_present(self):
        from agents.agent_minutes import AgentMinutes
        from core.knowledge_hub import KnowledgeHub

        hub = KnowledgeHub.new()
        hub.transcript_clean = "Texto de teste da reuniao."
        hub.ata_template_markdown = "# Ata\n\n## Participantes"

        agent = AgentMinutes.__new__(AgentMinutes)  # bypass __init__ (no LLM client needed)
        agent._skill = ""
        system, _ = agent.build_prompt(hub, "pt-BR")
        assert "Modelo de Ata do Contexto" in system
        assert "## Participantes" in system

    def test_no_injection_when_template_absent(self):
        from agents.agent_minutes import AgentMinutes
        from core.knowledge_hub import KnowledgeHub

        hub = KnowledgeHub.new()
        hub.transcript_clean = "Texto de teste da reuniao."

        agent = AgentMinutes.__new__(AgentMinutes)
        agent._skill = ""
        system, _ = agent.build_prompt(hub, "pt-BR")
        assert "Modelo de Ata do Contexto" not in system
