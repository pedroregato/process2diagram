# modules/minutes_exporter.py
# ─────────────────────────────────────────────────────────────────────────────
# Export MinutesModel to Word (.docx), PDF, and standalone HTML.
#
# to_docx(minutes) → bytes  — python-docx
# to_pdf(minutes)  → bytes  — fpdf2 (pure Python, Latin-1 covers Portuguese)
# to_html(minutes) → str    — self-contained HTML, no CDN dependency; falls
#                              back to parsing minutes_md when structured
#                              fields are empty (meeting loaded from DB)
# ─────────────────────────────────────────────────────────────────────────────

from __future__ import annotations

from datetime import datetime
from io import BytesIO
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from core.knowledge_hub import MinutesModel


# ── Shared helpers ────────────────────────────────────────────────────────────

def _now() -> str:
    return datetime.now().strftime("%d/%m/%Y")


def _prio_label(p: str) -> str:
    return {"high": "Alta", "normal": "Normal", "low": "Baixa"}.get(p, p.capitalize())


# ── PC160 Fase 2 — template-shaped section rendering ────────────────────────
# Maps a reference .docx section (by heading text) to a MinutesModel field,
# and — when the reference showed that section as a table — to the table's
# columns, so the exported .docx follows the SAME visual shape (table vs.
# list vs. paragraph) as whatever template the context registered, without
# hardcoding any particular template's section names or columns. See
# modules/ata_template_engine.py::extract_template_from_docx() for how
# style_spec["sections"] is derived.

import re as _re


def _kw(*roots: str) -> "_re.Pattern":
    """Compiles '<root1>|<root2>|...' with a leading \\b on each root, so a
    short/generic root (e.g. 'tema', 'depend') only matches at the START of
    a word — never as a substring inside an unrelated word ('sistema',
    'independente'). No trailing \\b, so Portuguese inflections (plural,
    gender) after the root still match ('tema' -> 'temas'/'temática')."""
    return _re.compile("|".join(rf"\b{root}" for root in roots), _re.I)


# Longer/rarer roots first when a heading could plausibly contain two
# matches at once (e.g. "Assuntos Discutidos / Decisões Tomadas" contains
# both a 'summary' signal — assunto/discuti — and a 'decisions' one —
# decis); the more specific term wins.
_SECTION_FIELD_PATTERNS: list[tuple["_re.Pattern", str]] = [
    (_kw("participant", "presente"), "participants"),
    (_kw("decis"), "decisions"),
    (_kw("a[cç][aã]o.*respons", "encaminha", "item.*a[cç][aã]o", "next step", "tarefa"), "action_items"),
    (_kw("proxima reuni", "próxima reuni"), "next_meeting"),
    (_kw("risco"), "risks_identified"),
    (_kw("depend"), "dependencies"),
    (_kw("pergunta.*abert", "questao.*abert", "questão.*abert", "open question"), "open_questions"),
    (_kw("premissa", "assumption"), "assumptions"),
    (_kw("stakeholder", "parte interessada"), "stakeholder_needs"),
    (_kw("pauta", "agenda", "tema"), "agenda"),
    (_kw("resumo", "assunto", "discuti", "topico", "tópico"), "summary"),
]


def _resolve_section_field(heading_text: str) -> str | None:
    """First matching synonym pattern wins. None = no known MinutesModel
    field for this heading — the section is skipped entirely at render
    time rather than fabricating content for it."""
    for pattern, field in _SECTION_FIELD_PATTERNS:
        if pattern.search(heading_text or ""):
            return field
    return None


_COLUMN_FIELD_PATTERNS: dict[str, list[tuple["_re.Pattern", str | None]]] = {
    "participants": [
        (_kw("nome", "name"), None),  # None = the raw string itself
    ],
    "action_items": [
        (_kw("tarefa", "a[cç][aã]o", "atividade"), "task"),
        (_kw("respons"), "responsible"),
        (_kw("prazo", "data", "deadline", "due"), "deadline"),
        (_kw("priorid"), "priority"),
        (_kw("levantado", "solicitante", "por quem", "raised"), "raised_by"),
    ],
}


def _shade_table_header_row(table, accent_rgb_hex: str) -> None:
    """Fills the first row's cells with accent_rgb_hex and sets white bold
    text — the same visual treatment already used for the default
    action_items table, now shared across any templated table too."""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    if not table.rows:
        return
    for cell in table.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), str(accent_rgb_hex))
        tcPr.append(shd)


def _render_field_as_table(doc, field: str, minutes: "MinutesModel", columns: list[str], accent) -> None:
    """
    Draws a table with EXACTLY the columns captured from the reference
    doc. Each detected column is matched (via _COLUMN_FIELD_PATTERNS) to
    an attribute; unmatched columns keep their header but every cell is
    "—" (same convention already used for ai.deadline/ai.raised_by when
    absent) — never fabricates data the transcript didn't provide (e.g.
    participant e-mail/unit). Assumes columns is non-empty; callers
    degrade to list rendering when the reference table had no readable
    header row.
    """
    from docx.shared import Pt

    patterns = _COLUMN_FIELD_PATTERNS.get(field, [])
    col_attrs: list[str | None] = []
    for col_name in columns:
        attr = next((a for pat, a in patterns if pat.search(col_name)), "__unmatched__")
        col_attrs.append(attr)

    rows_data = getattr(minutes, field) or []

    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(columns):
        hdr_cells[i].text = col_name

    for item in rows_data:
        row_cells = table.add_row().cells
        for i, attr in enumerate(col_attrs):
            if attr == "__unmatched__":
                value = "—"
            elif attr is None:
                value = str(item)
            else:
                value = getattr(item, attr, None) or "—"
            row_cells[i].text = str(value)
        for cell in row_cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    _shade_table_header_row(table, str(accent))
    doc.add_paragraph()


def _render_templated_sections(doc, minutes: "MinutesModel", sections: list[dict], navy, accent) -> set[str]:
    """
    Walks `sections` (from style_spec, document order) and renders each
    recognized one following the SHAPE captured from the reference doc
    (table/list/paragraph). Returns the set of MinutesModel field names
    rendered here, so to_docx() can still fall back to its default
    rendering for any field with real data that had no matching section
    in the template (never silently drops content).

    Never fabricates: a section with no known field is skipped; a table
    column with no data source is rendered as "—" rather than invented;
    a heading that repeats (maps to an already-rendered field — e.g. a
    template with "Participantes" appearing once per meeting sub-block)
    is skipped the second time, since MinutesModel holds one flat list
    per field, already printed in full on the first occurrence.
    """
    from docx.shared import Pt

    rendered: set[str] = set()

    for section in sections:
        field = _resolve_section_field(section.get("heading_text", ""))
        if field is None or field in rendered:
            continue
        value = getattr(minutes, field, None)
        if not value:
            continue

        heading_label = {
            "participants": "Participantes",
            "agenda": "Pauta",
            "summary": "Resumo da Reunião",
            "decisions": "Decisões Tomadas",
            "action_items": "Encaminhamentos / Action Items",
            "next_meeting": "Próxima Reunião",
            "risks_identified": "Riscos Identificados",
            "dependencies": "Dependências",
            "open_questions": "Questões em Aberto",
            "assumptions": "Premissas",
            "stakeholder_needs": "Necessidades dos Stakeholders",
        }.get(field, field)

        content_kind = section.get("content_kind")
        columns = section.get("columns") or []

        if content_kind == "table" and columns and field in _COLUMN_FIELD_PATTERNS:
            _heading_with_border(doc, heading_label, accent)
            _render_field_as_table(doc, field, minutes, columns, accent)
        elif field == "action_items":
            _heading_with_border(doc, f"{heading_label} ({len(value)})", accent)
            _render_default_action_items_table(doc, minutes)
        elif field == "summary":
            _heading_with_border(doc, heading_label, accent)
            _render_default_summary(doc, minutes, navy)
        elif field == "next_meeting":
            _heading_with_border(doc, heading_label, accent)
            p = doc.add_paragraph()
            run = p.add_run(str(value))
            run.font.size = Pt(11)
            run.font.color.rgb = navy
            doc.add_paragraph()
        elif field == "agenda":
            _heading_with_border(doc, heading_label, accent)
            for i, item in enumerate(value, 1):
                p = doc.add_paragraph(style="List Number")
                p.add_run(item).font.size = Pt(11)
                p.paragraph_format.space_after = Pt(2)
            doc.add_paragraph()
        else:
            # participants / decisions / risks / dependencies / open_questions /
            # assumptions / stakeholder_needs — flat list[str] fields, bulleted.
            _heading_with_border(doc, heading_label, accent)
            for item in value:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(str(item)).font.size = Pt(11)
                p.paragraph_format.space_after = Pt(2)
            doc.add_paragraph()

        rendered.add(field)

    return rendered


def _heading_with_border(doc, text: str, accent) -> None:
    """Section heading with a bottom-border separator in the accent color
    — shared by to_docx()'s default rendering and _render_templated_sections()."""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    h = doc.add_paragraph()
    run = h.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = accent
    run.font.name = "Calibri"
    pPr = h._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), str(accent))
    pBdr.append(bottom)
    pPr.append(pBdr)
    h.paragraph_format.space_after = Pt(4)


def _render_default_action_items_table(doc, minutes: "MinutesModel") -> None:
    """The app's fixed 5-column action items table (Prioridade/Tarefa/
    Responsável/Prazo/Levantado por), header shaded navy — unchanged
    default shape, used both as the fallback when no template applies
    and when a template's 'action_items'-mapped section wasn't itself a
    recognizable table."""
    from docx.shared import Pt

    headers = ["Prioridade", "Tarefa", "Responsável", "Prazo", "Levantado por"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h

    for ai in minutes.action_items:
        row_cells = table.add_row().cells
        row_cells[0].text = _prio_label(ai.priority)
        row_cells[1].text = ai.task
        row_cells[2].text = ai.responsible
        row_cells[3].text = ai.deadline or "—"
        row_cells[4].text = ai.raised_by or "—"
        for cell in row_cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    _shade_table_header_row(table, "0B1E3D")
    doc.add_paragraph()


def _render_default_summary(doc, minutes: "MinutesModel", navy) -> None:
    from docx.shared import Pt

    for block in minutes.summary:
        topic = block.get("topic", "")
        content = block.get("content", "")
        if topic:
            sub = doc.add_paragraph()
            run = sub.add_run(topic)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = navy
            sub.paragraph_format.space_before = Pt(6)
            sub.paragraph_format.space_after = Pt(2)
        if content:
            p = doc.add_paragraph(content)
            p.paragraph_format.space_after = Pt(4)
    doc.add_paragraph()


def _render_markdown_docx(doc, md: str, navy, accent, *, page_break_before_h2: bool = False) -> None:
    """
    Minimal Markdown -> docx renderer, used only when a meeting was loaded
    from the DB without structured minutes fields (only minutes_md
    persisted) — mirrors to_html()'s _md_to_html_fallback so the Word
    export isn't a near-empty stub in that case. Handles: #/##/### headers,
    -/* bullets, 1. numbered lists, | table | rows, blank-line paragraph
    breaks, plain paragraphs. Inline **bold**/*italic* markers are
    stripped (not rendered as bold/italic runs — not worth per-token
    splitting for a fallback path).

    page_break_before_h2: opt-in (default off) — starts each ## section on
    a new page, for multi-section consolidated exports (e.g.
    exportar_pacote_completo) where sections read as separate chapters.
    """
    import re
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _section_border(paragraph, color) -> None:
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), str(color))
        pBdr.append(bottom)
        pPr.append(pBdr)
        paragraph.paragraph_format.space_after = Pt(4)

    def _strip_inline(text: str) -> str:
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        text = re.sub(r"\*(.+?)\*", r"\1", text)
        return text

    lines = (md or "").splitlines()
    table_rows: list[list[str]] = []
    _seen_h2 = False

    def _flush_table() -> None:
        nonlocal table_rows
        if not table_rows:
            return
        n_cols = max(len(r) for r in table_rows)
        table = doc.add_table(rows=0, cols=n_cols)
        table.style = "Table Grid"
        for row in table_rows:
            cells = table.add_row().cells
            for i in range(n_cols):
                cells[i].text = row[i] if i < len(row) else ""
                for para in cells[i].paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)
        doc.add_paragraph()
        table_rows = []

    for line in lines:
        stripped = line.rstrip()

        if stripped.startswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            if all(re.match(r"^[-:]+$", c) for c in cells if c):
                continue  # separator row
            table_rows.append([_strip_inline(c) for c in cells])
            continue
        _flush_table()

        if stripped.startswith("### "):
            p = doc.add_paragraph()
            r = p.add_run(_strip_inline(stripped[4:]))
            r.bold = True; r.font.size = Pt(12); r.font.color.rgb = accent
        elif stripped.startswith("## "):
            if page_break_before_h2 and _seen_h2:
                doc.add_page_break()
            _seen_h2 = True
            p = doc.add_paragraph()
            r = p.add_run(_strip_inline(stripped[3:]))
            r.bold = True; r.font.size = Pt(14); r.font.color.rgb = navy
            _section_border(p, accent)
        elif stripped.startswith("# "):
            p = doc.add_paragraph()
            r = p.add_run(_strip_inline(stripped[2:]))
            r.bold = True; r.font.size = Pt(16); r.font.color.rgb = navy
        elif re.match(r"^[-*] ", stripped):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(_strip_inline(stripped[2:])).font.size = Pt(11)
        elif re.match(r"^\d+\. ", stripped):
            p = doc.add_paragraph(style="List Number")
            p.add_run(_strip_inline(re.sub(r"^\d+\.\s*", "", stripped))).font.size = Pt(11)
        elif not stripped:
            continue
        else:
            p = doc.add_paragraph()
            p.add_run(_strip_inline(stripped)).font.size = Pt(11)

    _flush_table()


def _add_page_number_footer(doc) -> None:
    """Inserts a 'Página X de Y' PAGE/NUMPAGES field in every section's
    footer, centered. Standard python-docx raw-XML field-code technique —
    python-docx has no high-level API for dynamic fields (they're computed
    by Word on open/print, not baked into the file)."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _field(instr: str):
        run = OxmlElement("w:r")
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr_el = OxmlElement("w:instrText")
        instr_el.set(qn("xml:space"), "preserve")
        instr_el.text = instr
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        for el in (fld_begin,):
            run.append(el)
        run2 = OxmlElement("w:r")
        run2.append(instr_el)
        run3 = OxmlElement("w:r")
        run3.append(fld_sep)
        run4 = OxmlElement("w:r")
        run4.append(fld_end)
        return [run, run2, run3, run4]

    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("Página ")
        for el in _field("PAGE"):
            p._p.append(el)
        p.add_run(" de ")
        for el in _field("NUMPAGES"):
            p._p.append(el)


def _force_field_recalculation(doc) -> None:
    """Sets <w:updateFields w:val="true"/> in word/settings.xml — without
    this, Word/LibreOffice show PAGE/NUMPAGES fields as blank or "0" until
    the user manually selects all + F9 ("Update Field"), because raw
    XML-inserted fields (as opposed to ones Word itself inserted via the
    UI) aren't marked dirty/stale on their own. This is the actual fix for
    page numbers silently not appearing on first open."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    settings = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


def _add_document_header(doc, title: str, subtitle: str = "") -> None:
    """Adds a running header (repeats on every page) with the project/
    document identification — distinct from the body's H1 title, which
    only appears once on page 1. Standard convention for multi-page
    corporate reports."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    ACCENT = RGBColor(0x2E, 0x7F, 0xD9)
    MUTED  = RGBColor(0x64, 0x74, 0x8B)

    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(title)
        r.bold = True; r.font.size = Pt(10); r.font.color.rgb = ACCENT
        if subtitle:
            r2 = p.add_run(f"  ·  {subtitle}")
            r2.font.size = Pt(9); r2.font.color.rgb = MUTED


def markdown_to_docx(
    md_text: str,
    *,
    add_page_numbers: bool = False,
    header_title: str | None = None,
    header_subtitle: str = "",
    page_break_before_h2: bool = False,
) -> bytes:
    """
    Generic Markdown -> .docx converter, for LLM-generated artifacts that
    aren't a MinutesModel (Project Charter, release notes, etc.) — reuses
    _render_markdown_docx() (same headers/bullets/numbered-lists/tables
    support already exercised by the ata export fallback) so there's a
    single Markdown->docx implementation in the app, not two.

    All extras are opt-in (default off/None), so output bytes for existing
    callers (export_project_charter_docx) don't change unless requested:
      add_page_numbers    — centered "Página X de Y" footer field, forced
                             to recalculate on open (see
                             _force_field_recalculation).
      header_title         — running header repeated on every page (e.g.
                             the project/context name) — distinct from the
                             body's H1, which only shows on page 1.
      header_subtitle       — smaller, muted text next to header_title
                             (e.g. document type + date).
      page_break_before_h2 — starts each ## section on its own page.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm

    NAVY   = RGBColor(0x0B, 0x1E, 0x3D)
    ACCENT = RGBColor(0x2E, 0x7F, 0xD9)

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    _render_markdown_docx(doc, md_text, NAVY, ACCENT, page_break_before_h2=page_break_before_h2)

    if header_title:
        _add_document_header(doc, header_title, header_subtitle)

    if add_page_numbers:
        _add_page_number_footer(doc)
        _force_field_recalculation(doc)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Word (.docx) ──────────────────────────────────────────────────────────────

def to_docx(minutes: "MinutesModel", template_spec: dict | None = None) -> bytes:
    """
    Generate a Word document from MinutesModel. Returns raw bytes.

    template_spec (PC160, optional — retrocompatible, omit for the default
    layout): {"accent_color": "#RRGGBB" | None, "assets": [{"asset_type",
    "image_bytes", ...}, ...]}, as produced by
    modules/ata_template_engine.py::extract_template_from_docx(). When
    present:
      - "accent_color" overrides the section-heading color (ACCENT below);
        the base NAVY brand color is intentionally left untouched — Fase 1
        of the template feature only carries one color, per
        melhorias/templates-ata-por-contexto.md.
      - the first asset with asset_type in ("logo", "header_image") is
        inserted into the document header;
      - "sections" (PC160 Fase 2, optional) — per-section shape captured
        from the reference .docx ("table"/"list"/"paragraph" + columns,
        see modules/ata_template_engine.py::extract_template_from_docx())
        — recognized sections render in that shape instead of the fixed
        default layout below; unrecognized sections are skipped rather
        than fabricated, and any field with real data but no matching
        section still falls back to the default rendering.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    NAVY   = RGBColor(0x0B, 0x1E, 0x3D)
    ACCENT = RGBColor(0x2E, 0x7F, 0xD9)
    MUTED  = RGBColor(0x64, 0x74, 0x8B)

    _accent_hex = (template_spec or {}).get("accent_color")
    if _accent_hex:
        try:
            _h = _accent_hex.lstrip("#")
            ACCENT = RGBColor(int(_h[0:2], 16), int(_h[2:4], 16), int(_h[4:6], 16))
        except (ValueError, IndexError):
            pass  # malformed color — keep the app default rather than fail the export

    doc = Document()

    # ── Page margins ─────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Template logo (header) ──────────────────────────────────────────────
    _assets = (template_spec or {}).get("assets") or []
    _logo = next(
        (a for a in _assets if a.get("asset_type") in ("logo", "header_image")), None
    )
    if _logo and _logo.get("image_bytes"):
        try:
            header = doc.sections[0].header
            header.is_linked_to_previous = False
            p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(BytesIO(_logo["image_bytes"]), height=Cm(1.5))
        except Exception:
            pass  # malformed/unsupported image — export continues without the logo

    # ── Default paragraph style ───────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Title ─────────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(minutes.title or "Ata de Reunião")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = NAVY

    # ── Metadata line ─────────────────────────────────────────────────────────
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parts = []
    if minutes.date:
        parts.append(f"Data: {minutes.date}")
    if minutes.location:
        parts.append(f"Local/Modalidade: {minutes.location}")
    parts.append(f"Gerado em: {_now()}")
    for i, part in enumerate(parts):
        if i:
            r = meta.add_run("   ·   ")
            r.font.color.rgb = MUTED
        r = meta.add_run(part)
        r.font.color.rgb = MUTED
        r.font.size = Pt(10)

    doc.add_paragraph()  # spacer

    def _heading(text: str) -> None:
        _heading_with_border(doc, text, ACCENT)

    def _bullet(text: str) -> None:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(text).font.size = Pt(11)
        p.paragraph_format.space_after = Pt(2)

    def _body(text: str) -> None:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(4)

    # ── Markdown fallback ─────────────────────────────────────────────────────
    # Meeting loaded from DB without structured fields (only minutes_md
    # persisted) — render the raw markdown instead of an empty document.
    # The structured sections below are all no-ops in this case (their
    # `if minutes.xxx:` guards see empty lists), so no early return needed.
    _has_structured = bool(
        minutes.participants or minutes.agenda or minutes.summary
        or minutes.decisions or minutes.action_items
        or minutes.assumptions or minutes.open_questions
        or minutes.risks_identified or minutes.dependencies or minutes.stakeholder_needs
    )
    if not _has_structured and getattr(minutes, "minutes_md", ""):
        _render_markdown_docx(doc, minutes.minutes_md, NAVY, ACCENT)

    # ── PC160 Fase 2 — template-shaped sections (table/list/paragraph) ────────
    # Renders whatever the reference .docx recognizably mapped to a
    # MinutesModel field, following its captured shape. No-op (empty set)
    # when the template has no "sections" (older/plain templates, or none
    # active) — every block below then runs exactly as before.
    _sections = (template_spec or {}).get("sections") or []
    _rendered: set[str] = set()
    if _has_structured and _sections:
        _rendered = _render_templated_sections(doc, minutes, _sections, NAVY, ACCENT)

    # ── Participants ──────────────────────────────────────────────────────────
    if "participants" not in _rendered and minutes.participants:
        _heading("Participantes")
        for p in minutes.participants:
            _bullet(p)
        doc.add_paragraph()

    # ── Agenda ────────────────────────────────────────────────────────────────
    if "agenda" not in _rendered and minutes.agenda:
        _heading("Pauta")
        for i, item in enumerate(minutes.agenda, 1):
            p = doc.add_paragraph(style="List Number")
            p.add_run(item).font.size = Pt(11)
            p.paragraph_format.space_after = Pt(2)
        doc.add_paragraph()

    # ── Summary ───────────────────────────────────────────────────────────────
    if "summary" not in _rendered and minutes.summary:
        _heading("Resumo da Reunião")
        _render_default_summary(doc, minutes, NAVY)

    # ── Decisions ─────────────────────────────────────────────────────────────
    if "decisions" not in _rendered and minutes.decisions:
        _heading("Decisões Tomadas")
        for d in minutes.decisions:
            _bullet(d)
        doc.add_paragraph()

    # ── Action Items table ────────────────────────────────────────────────────
    if "action_items" not in _rendered and minutes.action_items:
        _heading(f"Encaminhamentos / Action Items ({len(minutes.action_items)})")
        doc.add_paragraph()
        _render_default_action_items_table(doc, minutes)

    # ── Next meeting ──────────────────────────────────────────────────────────
    if "next_meeting" not in _rendered and minutes.next_meeting:
        _heading("Próxima Reunião")
        p = doc.add_paragraph()
        run = p.add_run(minutes.next_meeting)
        run.font.size = Pt(11)
        run.font.color.rgb = NAVY
        doc.add_paragraph()

    # ── BABOK fields (assumptions/open_questions/risks/dependencies/
    # stakeholder_needs) — bulleted, same shape as decisions/participants.
    # Previously never rendered in the .docx export at all; now shares the
    # same _rendered-guard as every other field so a template section can
    # claim any of them, and each still falls back here otherwise.
    for _field, _label in (
        ("assumptions", "Premissas"),
        ("open_questions", "Questões em Aberto"),
        ("risks_identified", "Riscos Identificados"),
        ("dependencies", "Dependências"),
        ("stakeholder_needs", "Necessidades dos Stakeholders"),
    ):
        _value = getattr(minutes, _field, None)
        if _field not in _rendered and _value:
            _heading(_label)
            for _item in _value:
                _bullet(_item)
            doc.add_paragraph()

    # ── Footer note ───────────────────────────────────────────────────────────
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run(f"Documento gerado automaticamente por Process2Diagram — {_now()}")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    run.italic = True

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── PDF ───────────────────────────────────────────────────────────────────────

# fpdf2 core fonts (Helvetica) are limited to Latin-1 (ISO 8859-1).
# This helper replaces common non-Latin-1 characters with safe equivalents
# before passing text to fpdf cells/multi_cell calls.
_PDF_CHAR_MAP = {
    "\u2014": "-",    # em dash
    "\u2013": "-",    # en dash
    "\u2012": "-",    # figure dash
    "\u2010": "-",    # hyphen
    "\u2022": "*",    # bullet
    "\u2023": ">",    # triangular bullet
    "\u2192": "->",   # right arrow
    "\u2190": "<-",   # left arrow
    "\u2026": "...",  # horizontal ellipsis
    "\u201c": '"',    # left double quotation mark
    "\u201d": '"',    # right double quotation mark
    "\u2018": "'",    # left single quotation mark
    "\u2019": "'",    # right single quotation mark
    "\u00b7": "*",    # middle dot
}


def _p(text: str) -> str:
    """Sanitize text for fpdf2 Latin-1 core fonts."""
    for src, dst in _PDF_CHAR_MAP.items():
        text = text.replace(src, dst)
    return text.encode("latin-1", errors="replace").decode("latin-1")


def to_pdf(minutes: "MinutesModel") -> bytes:
    """Generate a PDF from MinutesModel using fpdf2. Returns raw bytes."""
    from fpdf import FPDF

    class _PDF(FPDF):
        def header(self):
            pass  # custom header handled manually

        def footer(self):
            self.set_y(-14)
            self.set_font("Helvetica", "I", 8)
            self.set_text_color(100, 116, 139)
            self.cell(0, 8,
                _p(f"Process2Diagram  {_now()}  |  Pagina {self.page_no()}"),
                align="C")

    pdf = _PDF(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.set_margins(left=20, top=20, right=20)
    pdf.add_page()

    W = pdf.w - 40  # usable width (margins 20 each side)

    def set_navy():   pdf.set_text_color(11, 30, 61)
    def set_accent(): pdf.set_text_color(46, 127, 217)
    def set_muted():  pdf.set_text_color(100, 116, 139)
    def set_white():  pdf.set_text_color(255, 255, 255)
    def set_body():   pdf.set_text_color(28, 42, 58)

    # ── Title block ───────────────────────────────────────────────────────────
    pdf.set_fill_color(11, 30, 61)
    pdf.rect(20, 20, W, 24, style="F")
    pdf.set_xy(20, 22)
    pdf.set_font("Helvetica", "B", 16)
    set_white()
    title_text = _p(minutes.title or "Ata de Reuniao")
    pdf.multi_cell(W, 8, title_text, align="C")

    # Metadata strip
    pdf.set_fill_color(238, 244, 252)
    meta_y = pdf.get_y()
    pdf.rect(20, meta_y, W, 8, style="F")
    pdf.set_xy(20, meta_y + 1)
    pdf.set_font("Helvetica", "", 9)
    set_muted()
    meta_parts = []
    if minutes.date:
        meta_parts.append(f"Data: {_p(minutes.date)}")
    if minutes.location:
        meta_parts.append(f"Local: {_p(minutes.location)}")
    meta_parts.append(f"Gerado em: {_now()}")
    pdf.cell(W, 6, "   |   ".join(meta_parts), align="C")
    pdf.ln(12)

    def section_header(title: str) -> None:
        pdf.set_fill_color(46, 127, 217)
        pdf.set_font("Helvetica", "B", 9)
        set_white()
        pdf.cell(W, 7, _p(f"  {title.upper()}"), fill=True, ln=True)
        pdf.ln(2)
        set_body()
        pdf.set_font("Helvetica", "", 10)

    def bullet(text: str) -> None:
        pdf.set_font("Helvetica", "", 10)
        set_body()
        pdf.set_x(24)
        pdf.cell(4, 5, "-")
        pdf.set_x(28)
        pdf.multi_cell(W - 8, 5, _p(text))

    def body_text(text: str) -> None:
        pdf.set_font("Helvetica", "", 10)
        set_body()
        pdf.multi_cell(W, 5, _p(text))
        pdf.ln(1)

    # ── Participants ──────────────────────────────────────────────────────────
    if minutes.participants:
        section_header(f"Participantes ({len(minutes.participants)})")
        # Two-column layout
        col_w = W / 2 - 4
        for i, p in enumerate(minutes.participants):
            if i % 2 == 0:
                x_start = 20
            else:
                x_start = 20 + col_w + 8
            pdf.set_x(x_start)
            pdf.set_font("Helvetica", "", 10)
            set_body()
            pdf.cell(col_w, 5, _p(f"*  {p[:55]}"))
            if i % 2 == 1 or i == len(minutes.participants) - 1:
                pdf.ln(5)
        pdf.ln(4)

    # ── Agenda ────────────────────────────────────────────────────────────────
    if minutes.agenda:
        section_header(f"Pauta ({len(minutes.agenda)} itens)")
        for i, item in enumerate(minutes.agenda, 1):
            pdf.set_font("Helvetica", "", 10)
            set_body()
            pdf.set_x(24)
            pdf.cell(6, 5, f"{i}.")
            pdf.set_x(30)
            pdf.multi_cell(W - 10, 5, _p(item))
        pdf.ln(4)

    # ── Summary ───────────────────────────────────────────────────────────────
    if minutes.summary:
        section_header("Resumo da Reuniao")
        for block in minutes.summary:
            topic   = block.get("topic", "")
            content = block.get("content", "")
            if topic:
                pdf.set_font("Helvetica", "B", 10)
                set_navy()
                pdf.multi_cell(W, 5, _p(topic))
                pdf.ln(1)
            if content:
                body_text(content)
        pdf.ln(2)

    # ── Decisions ─────────────────────────────────────────────────────────────
    if minutes.decisions:
        section_header(f"Decisoes Tomadas ({len(minutes.decisions)})")
        for d in minutes.decisions:
            bullet(d)
        pdf.ln(4)

    # ── Action Items table ────────────────────────────────────────────────────
    if minutes.action_items:
        section_header(f"Encaminhamentos / Action Items ({len(minutes.action_items)})")
        col_widths = [22, W - 22 - 38 - 24 - 22, 38, 24, 22]
        headers    = ["Prior.", "Tarefa", "Responsavel", "Prazo", "Por"]

        # Table header
        pdf.set_fill_color(11, 30, 61)
        pdf.set_font("Helvetica", "B", 9)
        set_white()
        for h, w in zip(headers, col_widths):
            pdf.cell(w, 7, f" {h}", fill=True, border=0)
        pdf.ln()

        prio_colors = {
            "high":        (201, 123, 26),
            "normal":      (46,  127, 217),
            "low":         (26,  127, 90),
            "unspecified": (100, 116, 139),
        }

        for i, ai in enumerate(minutes.action_items):
            fill = i % 2 == 0
            pdf.set_fill_color(248, 250, 252) if fill else pdf.set_fill_color(255, 255, 255)
            pdf.set_font("Helvetica", "", 9)
            r, g, b = prio_colors.get(ai.priority, (100, 116, 139))
            pdf.set_text_color(r, g, b)
            pdf.cell(col_widths[0], 6, _p(f" {_prio_label(ai.priority)}"), fill=fill, border=0)
            set_body()
            # Task may be long — truncate for table
            task = ai.task if len(ai.task) <= 60 else ai.task[:57] + "..."
            pdf.cell(col_widths[1], 6, _p(f" {task}"), fill=fill, border=0)
            pdf.cell(col_widths[2], 6, _p(f" {ai.responsible[:20]}"), fill=fill, border=0)
            pdf.cell(col_widths[3], 6, _p(f" {ai.deadline or '-'}"), fill=fill, border=0)
            pdf.cell(col_widths[4], 6, _p(f" {ai.raised_by or '-'}"), fill=fill, border=0)
            pdf.ln()

        # Bottom border
        pdf.set_draw_color(213, 227, 245)
        pdf.line(20, pdf.get_y(), 20 + W, pdf.get_y())
        pdf.ln(5)

    # ── Next meeting ──────────────────────────────────────────────────────────
    if minutes.next_meeting:
        section_header("Proxima Reuniao")
        pdf.set_font("Helvetica", "B", 11)
        set_navy()
        pdf.cell(W, 7, _p(minutes.next_meeting), ln=True)
        pdf.ln(3)

    return bytes(pdf.output())


# ── HTML ──────────────────────────────────────────────────────────────────────
#
# Self-contained HTML (inline CSS, no CDN dependency) so the exported file
# opens correctly offline / via file://. Always available: renders the
# structured fields when present, or falls back to parsing minutes_md when
# a meeting was loaded from the DB with only the raw markdown persisted
# (structured fields — participants/agenda/summary/decisions/action_items —
# are not stored as separate columns, only minutes_md is).

_HTML_CSS = """
  * { box-sizing: border-box; }
  body {
    margin: 0; padding: 0;
    background: #F4F7FB;
    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Arial, sans-serif;
    color: #1e293b;
  }
  .page { max-width: 820px; margin: 0 auto; padding: 32px 24px 56px; }
  .card {
    background: #fff;
    border-radius: 10px;
    box-shadow: 0 1px 3px rgba(11,30,61,0.08);
    padding: 28px 32px;
    margin-bottom: 20px;
  }
  .title { text-align: center; color: #0B1E3D; font-size: 26px; font-weight: 700; margin: 0 0 8px; }
  .meta { text-align: center; color: #64748B; font-size: 13px; margin: 0; }
  .meta span + span::before { content: " · "; }
  h2.section-title {
    text-transform: uppercase;
    font-size: 12px; letter-spacing: 0.04em; font-weight: 700;
    color: #2E7FD9;
    border-bottom: 1px solid #2E7FD9;
    padding-bottom: 6px;
    margin: 0 0 12px;
  }
  ul.plain { margin: 0; padding-left: 20px; }
  ul.plain li { margin-bottom: 4px; }
  ol.plain { margin: 0; padding-left: 20px; }
  ol.plain li { margin-bottom: 4px; }
  .topic { color: #0B1E3D; font-weight: 700; margin: 14px 0 2px; }
  .topic:first-child { margin-top: 0; }
  p.body-text { margin: 0 0 10px; line-height: 1.5; }
  table.actions { width: 100%; border-collapse: collapse; font-size: 13px; }
  table.actions th {
    background: #0B1E3D; color: #fff; text-align: left;
    padding: 8px 10px; font-weight: 600;
  }
  table.actions td { padding: 8px 10px; border-bottom: 1px solid #E2E8F0; }
  table.actions tr:last-child td { border-bottom: none; }
  .prio { display: inline-block; padding: 2px 8px; border-radius: 999px; font-size: 11px; font-weight: 600; }
  .prio-high   { background: #FDECD8; color: #C9791A; }
  .prio-normal { background: #DCEAFB; color: #2E7FD9; }
  .prio-low    { background: #DBF0E6; color: #1A7F5A; }
  .next-meeting { color: #0B1E3D; font-size: 14px; }
  .footer { text-align: center; color: #64748B; font-size: 11px; font-style: italic; margin-top: 8px; }
  /* Markdown-fallback rendering (minutes_md-only, no structured fields) */
  .min-h3 { color: #0B1E3D; font-size: 18px; margin: 18px 0 8px; }
  .min-h4 { color: #0B1E3D; font-size: 15px; margin: 14px 0 6px; }
  .min-h5 { color: #2E7FD9; font-size: 13px; margin: 10px 0 4px; }
  table.minutes-table { width: 100%; border-collapse: collapse; font-size: 13px; margin: 8px 0 14px; }
  table.minutes-table td { padding: 6px 10px; border: 1px solid #E2E8F0; }
"""


def _md_to_html_fallback(md: str) -> str:
    """
    Minimal Markdown -> HTML converter, used only when a meeting was loaded
    from the DB without structured minutes fields (only minutes_md persisted).
    Handles: | table | rows, #/##/### headers, -/* bullets, 1. numbered
    lists, blank-line paragraph breaks, and inline **bold**/*italic*.
    """
    import html as _html_lib
    import re

    lines = (md or "").splitlines()
    out: list[str] = []
    in_ul = False
    in_table = False

    def close_ul() -> None:
        nonlocal in_ul
        if in_ul:
            out.append("</ul>")
            in_ul = False

    def close_table() -> None:
        nonlocal in_table
        if in_table:
            out.append("</tbody></table>")
            in_table = False

    def inline(text: str) -> str:
        text = _html_lib.escape(text)
        text = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", text)
        text = re.sub(r"\*(.+?)\*", r"<em>\1</em>", text)
        return text

    for line in lines:
        stripped = line.rstrip()

        if stripped.startswith("|"):
            close_ul()
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            if all(re.match(r"^[-:]+$", c) for c in cells if c):
                continue
            if not in_table:
                out.append('<table class="minutes-table"><tbody>')
                in_table = True
            tds = "".join(f"<td>{inline(c)}</td>" for c in cells)
            out.append(f"<tr>{tds}</tr>")
            continue

        close_table()

        if stripped.startswith("### "):
            close_ul()
            out.append(f'<h5 class="min-h5">{inline(stripped[4:])}</h5>')
        elif stripped.startswith("## "):
            close_ul()
            out.append(f'<h4 class="min-h4">{inline(stripped[3:])}</h4>')
        elif stripped.startswith("# "):
            close_ul()
            out.append(f'<h3 class="min-h3">{inline(stripped[2:])}</h3>')
        elif re.match(r"^[-*] ", stripped):
            if not in_ul:
                out.append('<ul class="plain">')
                in_ul = True
            out.append(f"<li>{inline(stripped[2:])}</li>")
        elif re.match(r"^\d+\. ", stripped):
            close_ul()
            out.append(f"<li>{inline(re.sub(r'^\\d+\\. ', '', stripped))}</li>")
        elif not stripped:
            close_ul()
        else:
            close_ul()
            out.append(f'<p class="body-text">{inline(stripped)}</p>')

    close_ul()
    close_table()
    return "\n".join(out)


def to_html(minutes: "MinutesModel") -> str:
    """
    Generate a self-contained HTML document from MinutesModel. Returns a
    full HTML string (not bytes) — always available regardless of whether
    structured fields (agenda/summary/decisions/action_items) are populated
    or the meeting only has raw minutes_md (loaded from DB without
    structured data). Same navy/accent palette as to_docx()/to_pdf().
    """
    import html as _html_lib

    def esc(s: object) -> str:
        return _html_lib.escape(str(s or ""))

    title = esc(minutes.title or "Ata de Reunião")

    meta_parts = []
    if minutes.date:
        meta_parts.append(f"<span>Data: {esc(minutes.date)}</span>")
    if minutes.location:
        meta_parts.append(f"<span>Local/Modalidade: {esc(minutes.location)}</span>")
    meta_parts.append(f"<span>Gerado em: {esc(_now())}</span>")
    meta_html = "".join(meta_parts)

    has_structured = bool(
        minutes.participants or minutes.agenda or minutes.summary
        or minutes.decisions or minutes.action_items
    )

    sections: list[str] = []

    if has_structured:
        if minutes.participants:
            items = "".join(f"<li>{esc(p)}</li>" for p in minutes.participants)
            sections.append(
                f'<div class="card"><h2 class="section-title">Participantes</h2>'
                f'<ul class="plain">{items}</ul></div>'
            )
        if minutes.agenda:
            items = "".join(f"<li>{esc(a)}</li>" for a in minutes.agenda)
            sections.append(
                f'<div class="card"><h2 class="section-title">Pauta</h2>'
                f'<ol class="plain">{items}</ol></div>'
            )
        if minutes.summary:
            blocks = []
            for block in minutes.summary:
                topic = esc(block.get("topic", ""))
                content = esc(block.get("content", ""))
                if topic:
                    blocks.append(f'<div class="topic">{topic}</div>')
                if content:
                    blocks.append(f'<p class="body-text">{content}</p>')
            sections.append(
                f'<div class="card"><h2 class="section-title">Resumo da Reunião</h2>'
                f'{"".join(blocks)}</div>'
            )
        if minutes.decisions:
            items = "".join(f"<li>{esc(d)}</li>" for d in minutes.decisions)
            sections.append(
                f'<div class="card"><h2 class="section-title">Decisões Tomadas</h2>'
                f'<ul class="plain">{items}</ul></div>'
            )
        if minutes.action_items:
            rows = []
            for ai in minutes.action_items:
                prio_class = {"high": "prio-high", "normal": "prio-normal", "low": "prio-low"}.get(
                    ai.priority, "prio-normal"
                )
                rows.append(
                    "<tr>"
                    f'<td><span class="prio {prio_class}">{esc(_prio_label(ai.priority))}</span></td>'
                    f"<td>{esc(ai.task)}</td>"
                    f"<td>{esc(ai.responsible)}</td>"
                    f'<td>{esc(ai.deadline or "—")}</td>'
                    f'<td>{esc(ai.raised_by or "—")}</td>'
                    "</tr>"
                )
            sections.append(
                f'<div class="card"><h2 class="section-title">'
                f"Encaminhamentos / Action Items ({len(minutes.action_items)})</h2>"
                '<table class="actions"><thead><tr>'
                "<th>Prioridade</th><th>Tarefa</th><th>Responsável</th>"
                "<th>Prazo</th><th>Levantado por</th>"
                f"</tr></thead><tbody>{''.join(rows)}</tbody></table></div>"
            )
        if minutes.next_meeting:
            sections.append(
                f'<div class="card"><h2 class="section-title">Próxima Reunião</h2>'
                f'<p class="next-meeting">{esc(minutes.next_meeting)}</p></div>'
            )
    else:
        # Loaded from DB without structured fields — parse the raw markdown.
        body = _md_to_html_fallback(getattr(minutes, "minutes_md", ""))
        sections.append(f'<div class="card">{body}</div>')

    body_html = "\n".join(sections)

    return f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{title}</title>
<style>{_HTML_CSS}</style>
</head>
<body>
<div class="page">
  <div class="card">
    <p class="title">{title}</p>
    <p class="meta">{meta_html}</p>
  </div>
  {body_html}
  <p class="footer">Documento gerado automaticamente por Process2Diagram — {esc(_now())}</p>
</div>
</body>
</html>"""
